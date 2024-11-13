import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Import alignment options

def json_to_docx(json_file_path, docx_file_path):
    # Load JSON data
    with open(json_file_path, 'r') as json_file:
        data = json.load(json_file)
    
    # Create a new Document
    doc = Document()
    
    # Iterate over each entry in the JSON data
    for entry in data['images'][:10]:  # Limit to first 10 entries
        image_path = '../rsicd/imgs/' + entry.get("filename", "")
        description = entry["sentences"][0]["raw"]
        img_id = str(entry["imgid"])

        # Add separator after each entry
        doc.add_paragraph("-" * 20)
        doc.add_paragraph(img_id)
        
        # Add image if path is valid
        try:
            image_paragraph = doc.add_paragraph()  # Create a new paragraph for the image
            run = image_paragraph.add_run()
            run.add_picture(image_path, width=Inches(4))  # Add image
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the image paragraph
        except Exception as e:
            print(f"Failed to add image {image_path}: {e}")

        # Add original description with bold label
        if description:
            original_desc = doc.add_paragraph()
            original_desc.add_run("Descrição original: ").bold = True
            doc.add_paragraph()
            doc.add_paragraph(description)
            
            # Add newline between descriptions
            doc.add_paragraph()
            
            # Add edited description with bold label and newline
            edited_desc = doc.add_paragraph()
            edited_desc.add_run("Descrição editada: ").bold = True
            doc.add_paragraph()
            doc.add_paragraph(description)
        
        # Insert a page break to start the next entry on a new page
        doc.add_page_break()
        
    # Save the document
    doc.save(docx_file_path)
    print(f"Document saved as {docx_file_path}")

# Example usage
json_to_docx("../rsicd/dataset_rsicd.json", "output.docx")

